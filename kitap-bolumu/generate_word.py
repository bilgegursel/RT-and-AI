#!/usr/bin/env python3
"""Generate full styled Word chapter (~14-15 pages) with numbered refs, tables, figures."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import os
import re

OUT_DIR = "/workspace/kitap-bolumu"
FIG_DIR = os.path.join(OUT_DIR, "figures")
os.makedirs(FIG_DIR, exist_ok=True)

REFS = [
    "High KP, Bradley SF, Gravenstein S, et al. Clinical practice guideline for the evaluation of fever and infection in older adult residents of long-term care facilities: 2008 update by the Infectious Diseases Society of America. Clin Infect Dis. 2009;48(2):149-171.",
    "Norman DC. Fever in the elderly. Clin Infect Dis. 2000;31(1):148-151.",
    "Juthani-Mehta M, Petrilli CM, Wright PW, et al. Infectious diseases in older adults of long-term care facilities: update on approach to diagnosis and management. J Am Geriatr Soc. 2020;68(3):478-492.",
    "Gordon AL, Goodman A, Davies SL, et al. Diagnosis and management of sepsis in the older adult. BMJ. 2023;382:e075585.",
    "Evans SS, Aspnes GE, Safford MM. Aging and options to halt declining immunity to virus infections. Front Immunol. 2021;12:681449.",
    "Alnakhli AM, Alshammari A, Alshammari SA, et al. Aging and infection: impact of immunosenescence and inflammaging in respiratory viral infections. Cell Physiol Biochem. 2026;62(1):1-15.",
    "Elhassan O, El-Sherif M, El-Mashad A, et al. Atypical presentations of acute infections in hospitalized older adults: the prevalence, predictors, and outcomes. Egypt J Geriatr Gerontol. 2021;8(2):204172.",
    "Jump RLP, Sullivan GJ, Hemmert KC, et al. Reliability of nonlocalizing signs and symptoms as indicators of the presence of infection in nursing home residents. Infect Control Hosp Epidemiol. 2022;43(8):961-968.",
    "Ticinesi A, Nouvenne A, Meschi T. Screening and treatment of asymptomatic bacteriuria in older adults: the case against. J Am Geriatr Soc. 2022;70(8):2192-2198.",
    "Evans L, Rhodes A, Alhazzani W, et al. Surviving Sepsis Campaign: international guidelines for management of sepsis and septic shock 2021. Crit Care Med. 2021;49(11):e1063-e1143.",
    "Scottish Antimicrobial Prescribing Group. Good practice recommendations for antimicrobial use in frail older people. 2021. https://www.sapg.scot",
    "Carr E, Keating P, Gallagher P, et al. Safety and tolerability of antimicrobial agents in the older patient. Drugs Aging. 2023;40(12):1065-1083.",
    "Wales Medicines Strategy Group. Polypharmacy in older people: a guide for healthcare professionals. 2023.",
    "Fabbri E, Tebano G, De Angelis A, et al. Sepsis in frail older adults: tailored antimicrobial stewardship and individualized care approach. Antibiotics. 2026;15(5):496.",
    "Muller L, Chan CT, Santamaria LP, et al. Impact of immunosenescence on vaccine immune responses and countermeasures. Vaccines. 2024;12(11):1289.",
    "Aliberti S, Cook GS, Babu BL, et al. Multisociety guidance for infection prevention and control in nursing homes. Infect Control Hosp Epidemiol. 2025;46(12):10252.",
]
CORNERSTONE = {1, 2}


def setup_styles(doc):
    sec = doc.sections[0]
    for attr, val in [("top_margin", 2.5), ("bottom_margin", 2.5), ("left_margin", 2.5), ("right_margin", 2.5)]:
        setattr(sec, attr, Cm(val))
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i, size in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Times New Roman"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
        h.font.size = Pt(size)


def shade_cell(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def P(doc, text, bold=False, indent=True, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    return p


def P_cite(doc, segments, indent=True):
    """segments: list of str or int (ref num)"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    for seg in segments:
        if isinstance(seg, int):
            r = p.add_run(f"[{seg}]")
            r.font.superscript = True
            r.font.size = Pt(10)
        else:
            r = p.add_run(seg)
            r.font.size = Pt(12)
        r.font.name = "Times New Roman"
    return p


def H(doc, text, level=2):
    doc.add_heading(text, level=level)


def caption(doc, num, title, kind="Tablo"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{kind} {num}. {title}")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade_cell(c, "1A476F")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if ri % 2:
                shade_cell(c, "EDF2F9")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def algo_box(doc, lines, fig_num, title):
    caption(doc, fig_num, title, "Şekil")
    tb = doc.add_table(1, 1)
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tb.rows[0].cells[0]
    shade_cell(cell, "F5F8FC")
    cell.text = ""
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_img(doc, path, fig_num, title, w=15):
    caption(doc, fig_num, title, "Şekil")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(w))
    doc.add_paragraph()


def make_flowchart(path, boxes, colors=None):
    if colors is None:
        colors = ["#1A476F", "#2E75B6", "#5B9BD5", "#70AD47", "#FFC000", "#ED7D31"]
    n = len(boxes)
    fig, ax = plt.subplots(figsize=(7.5, max(2.5, n * 0.55 + 0.8)))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, n + 0.5)
    ax.axis("off")
    y = n
    for i, (t1, t2) in enumerate(boxes):
        c = colors[i % len(colors)]
        ax.add_patch(FancyBboxPatch((0.8, y - 0.32), 8.4, 0.62, boxstyle="round,pad=0.06",
                                    facecolor=c, edgecolor="#333", lw=1.1))
        txt = t1 if not t2 else f"{t1}\n{t2}"
        ax.text(5, y, txt, ha="center", va="center", fontsize=8.5,
                color="white" if i < 3 else "#222", fontweight="bold")
        if i < n - 1:
            ax.annotate("", xy=(5, y - 0.48), xytext=(5, y - 0.78),
                        arrowprops=dict(arrowstyle="->", color="#333", lw=1.4))
        y -= 1.05
    plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close()


def build_figures():
    f = {}
    f["s1"] = f"{FIG_DIR}/s1.png"
    make_flowchart(f["s1"], [
        ("Yaşlanma + komorbidite + kırılganlık", ""),
        ("İmmünosenesans + bariyer bozukluğu", ""),
        ("Atipik sunum + geç tanı", ""),
        ("Gecikmiş tedavi + organ yetmezliği", ""),
        ("Yüksek mortalite / geriatrik gerileme", ""),
    ])
    f["s2"] = f"{FIG_DIR}/s2.png"
    make_flowchart(f["s2"], [
        ("Yaşlı hasta – enfeksiyon şüphesi", ""),
        ("Septik şok → Stabilizasyon + antibiyotik ≤1 sa", ""),
        ("Orta-ağır → Yatış + tedavi ≤3 sa", ""),
        ("Stabil → Ayaktan + odak taraması", ""),
    ])
    f["s3"] = f"{FIG_DIR}/s3.png"
    make_flowchart(f["s3"], [
        ("Enfeksiyon olasılığı", ""),
        ("DÜŞÜK → İzlem, antibiyotik yok", ""),
        ("ORTA → Kültür + dar spektrum + komorbidite filtresi", ""),
        ("YÜKSEK/SEPSİS → Geniş spektrum + de-eskale", ""),
    ])
    f["s4"] = f"{FIG_DIR}/s4.png"
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    steps = [(5, 5.2, "Empirik antibiyotik seçimi"), (5, 4.1, "Acil tehlike / kaçınılacaklar"),
             (5, 3.0, "Odak-spesifik tercih"), (5, 1.9, "Komorbidite modifikasyonu"), (5, 0.8, "Rota, süre, de-eskale")]
    cols = ["#1A476F", "#2E75B6", "#5B9BD5", "#70AD47", "#FFC000"]
    for i, (x, y, txt) in enumerate(steps):
        ax.add_patch(FancyBboxPatch((x - 3.4, y - 0.3), 6.8, 0.55, boxstyle="round", facecolor=cols[i], edgecolor="#333"))
        ax.text(x, y, txt, ha="center", va="center", fontsize=9, color="white", fontweight="bold")
        if i < len(steps) - 1:
            ax.annotate("", xy=(x, steps[i + 1][1] + 0.3), xytext=(x, y - 0.3),
                        arrowprops=dict(arrowstyle="->", color="#333", lw=1.5))
    f["s4"] = f"{FIG_DIR}/s4.png"
    plt.savefig(f["s4"], dpi=200, bbox_inches="tight", facecolor="white")
    plt.close()
    f["s5"] = f"{FIG_DIR}/s5.png"
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.axis("off")
    ax.text(5, 4.5, "Sepsis – sıvı stratejisi (komorbiditeye göre)", ha="center", fontsize=11, fontweight="bold", color="#1A476F")
    ax.add_patch(FancyBboxPatch((3.2, 3.4), 3.6, 0.55, boxstyle="round", facecolor="#1A476F"))
    ax.text(5, 3.67, "SEPSİS / SEPTİK ŞOK", ha="center", va="center", color="white", fontsize=9, fontweight="bold")
    for x, txt, c in [(2.3, "KY yok\n500 mL", "#5B9BD5"), (5, "KY var\n250 mL + vazopressör", "#C00000"), (7.7, "KBH/diyaliz\nEkip koordinasyonu", "#7030A0")]:
        ax.annotate("", xy=(x, 2.6), xytext=(5, 3.4), arrowprops=dict(arrowstyle="->", color="#333", lw=1.2))
        ax.add_patch(FancyBboxPatch((x - 1.2, 1.8), 2.4, 0.9, boxstyle="round", facecolor=c))
        ax.text(x, 2.25, txt, ha="center", va="center", fontsize=8, color="white", fontweight="bold")
    plt.savefig(f["s5"], dpi=200, bbox_inches="tight", facecolor="white")
    plt.close()
    f["s6"] = f"{FIG_DIR}/s6.png"
    make_flowchart(f["s6"], [
        ("Enfeksiyon odağı belirlendi", ""),
        ("ÜSYE viral → Semptomatik (antibiyotik genellikle gerekmez)", ""),
        ("Pnömoni → Amoksisilin veya β-laktam (kinolon: KY/KBH/QT'de değiştir)", ""),
        ("İYE → TMP-SMX/nitrofurantoin/fosfomisin (eGFR'ye göre)", ""),
        ("Sepsis → Meropenem/pip-tazo + vankomisin; kaynak kontrolü", ""),
    ], ["#1A476F", "#A9D18E", "#2E75B6", "#5B9BD5", "#C00000"])
    return f


def build():
    figs = build_figures()
    doc = Document()
    setup_styles(doc)

    t = doc.add_heading("YAŞLILARDA ENFEKSİYON HASTALIKLARINA GENEL YAKLAŞIM", 1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    H(doc, "GİRİŞ")
    P_cite(doc, [
        "Enfeksiyon hastalıkları, yaşlı bireylerde morbidite ve mortalitenin en önemli nedenlerinden biridir. 65 yaş üstü popülasyonda solunum yolu enfeksiyonları, idrar yolu enfeksiyonları (İYE), deri-yumuşak doku enfeksiyonları, sepsis ve Clostridioides difficile enfeksiyonu sık görülür; hastane yatışları, fonksiyonel gerileme, deliryum ve uzun dönem bakım kuruluşu (UDBK) yerleşimi ile ilişkilidir. Genç yetişkinlerde geçerli olan “ateş + lokalizasyon bulgusu = enfeksiyon” modeli yaşlı hastada çoğu zaman yetersiz kalır",
        1, ". Bu bölümde, diğer kitap bölümlerinde olduğu gibi (hiperlipidemi, tiroid hastalıkları vb.), neden–niçin mantığı ve algoritmik yaklaşım üzerinden yaşlılarda enfeksiyon yönetiminin genel çerçevesi sunulmaktadır."
    ])
    P(doc, "Temel mesaj: Yaşlıda enfeksiyon yönetimi, tek bir organ odaklı değil; bütüncül geriatrik değerlendirme, komorbidite etkileşimi, kırılganlık (frailty) düzeyi ve tedavi hedeflerinin bireyselleştirilmesi üzerine kurulmalıdır.", bold=True)

    H(doc, "EPİDEMİYOLOJİ VE KLİNİK YÜK")
    P_cite(doc, [
        "Yaşlanan dünya nüfusunda enfeksiyonlar, kronik hastalıkların alevlenmesi ve akut geriatrik sendromların (deliryum, düşme, immobilite, inkontinans) önemli tetikleyicilerindendir. 75 yaş üstünde antibiyotik kullanım sıklığı genç yetişkinlere kıyasla yaklaşık iki kat; 90 yaş üstünde ve UDBK sakinlerinde ise üç kata kadar artmaktadır",
        3, ". Sepsis acil servis başvurularının önemli bir bölümünü oluşturur; sepsis nedeniyle acil yatışların ortalama yaşı 71’dir ",
        4, ". Enfeksiyon kaynaklı mortalite yaşlılarda belirgin şekilde yüksektir; bunun başlıca nedenleri gecikmiş tanı, atipik sunum, komorbid organ disfonksiyonu, immün yanıtın zayıflaması ve geç başlanan uygun antimikrobiyal tedavidir."
    ])

    caption(doc, 1, "Yaşlılarda sık görülen enfeksiyon sendromları ve tipik kaynaklar")
    table(doc, ["Sendrom", "Sık etkenler", "Yaşlıya özgü risk faktörleri"], [
        ["Solunum yolu enfeksiyonu", "S. pneumoniae, influenza, SARS-CoV-2, RSV", "KOAH, kalp yetmezliği, aşı eksikliği, aspirasyon"],
        ["İYE / piyelonefrit", "E. coli, Proteus, Klebsiella", "İdrar retansiyonu, kateter, diyabet"],
        ["Deri-yumuşak doku", "MRSA, streptokok, anaeroblar", "Dekubitus, venöz yetmezlik, yetersiz bakım"],
        ["Sepsis", "Gram negatif, MRSA, polimikrobiyal", "Kırılganlık, kateter, malnütrisyon"],
        ["C. difficile", "Toksin üreten C. difficile", "Antibiyotik maruziyeti, PPI, hospitalizasyon"],
    ], [4.5, 5, 5.5])

    H(doc, "YAŞLILARDA ENFEKSİYON RİSKİNİN NEDENLERİ: PATOFİZİYOLOJİK ÇERÇEVE")
    P(doc, "Yaşlıda artmış enfeksiyon riski tek bir mekanizmaya indirgenemez. Klinik karar verme için aşağıdaki başlıkların birlikte değerlendirilmesi gerekir.")

    H(doc, "1. İmmünosenesans ve inflamaging", 3)
    P_cite(doc, [
        "İmmünosenesans, yaşla birlikte hem doğuştan hem edinsel immün yanıtın değişmesidir. Timik involutu, naif T hücre rezervuarının azalması, bellek hücrelerinin artması, B hücre antikor yanıtının zayıflaması ve NK hücre aktivitesinde bozulma enfeksiyonlara karşı koruyuculuğu azaltır. İnflamaging ise kronik, düşük dereceli proinflamatuar durumdur (IL-6, TNF-α, CRP artışı); enfeksiyon sırasında aşırı veya yetersiz yanıt riskini artırır",
        5, 6, ". Klinik sonuç: Aşı yanıtı zayıflar; latent virüsler reaktive olabilir; viral ve bakteriyel enfeksiyonlar daha ağır seyreder."
    ])

    H(doc, "2. Yapısal ve fonksiyonel savunma bariyerlerinin bozulması", 3)
    for item in [
        "Solunum: Öksürük refleksinin zayıflaması, yutma disfonksiyonu (aspirasyon), azalmış vital kapasite.",
        "Deri: İncelme, kuruma, perfüzyon bozukluğu → mikrotravma ve kolonizasyon.",
        "Üriner sistem: Mesane boşaltım bozukluğu, prostat hipertrofisi, nörogen mesane, kalıcı kateter.",
        "GIS: Gastrik asit sekresyonunun azalması, bağırsak mikrobiyota değişimi.",
    ]:
        P(doc, "• " + item, indent=False)

    H(doc, "3. Komorbid hastalıklar ve iatrogenik faktörler", 3)
    P(doc, "Diyabet, kronik böbrek hastalığı (KBH), kalp yetmezliği, KOAH, demans, Parkinson hastalığı, malignite ve immünosupresif tedaviler enfeksiyon riskini artırır. Kateterler, cerrahi implantlar, yara pansumanı eksikliği ve UDBK yoğunluğu ek risk oluşturur.")

    H(doc, "4. Kırılganlık (frailty) ve geriatrik sendromlar", 3)
    P(doc, "Kırılganlık, stresörle (enfeksiyon gibi) karşılaşıldığında homeostazı koruyamama durumudur. Kırılgan yaşlıda enfeksiyon çoğu zaman deliryum, düşme, iştahsızlık, fonksiyonel gerileme ile manifest olur; klasik lokalizasyon bulgusu geç ortaya çıkabilir veya hiç görülmeyebilir.")

    H(doc, "5. Farmakokinetik/farmakodinamik değişiklikler", 3)
    P_cite(doc, [
        "Yaşlıda yağ dokusu artışı, total vücut suyu azalması, karaciğer ve böbrek klirensinde azalma nedeniyle antimikrobiyal dozaj hataları sık görülür. Polifarmasi ile ilaç–ilaç etkileşimleri advers olay riskini yükseltir",
        12, 13, "."
    ])

    add_img(doc, figs["s1"], 1, "Yaşlıda enfeksiyon riski: neden–sonuç zinciri")

    H(doc, "KLİNİK SUNUM: ATİPİK PREZENTASYON VE TANISAL TUZAKLAR")
    P(doc, "Yaşlı hastada enfeksiyon, genç hastadaki klasik tablodan farklı seyredebilir. Bu farkı bilmek, gereksiz antibiyotik kullanımını ve kaçırılan ciddi enfeksiyonları aynı anda azaltır.")

    H(doc, "Ateş yanıtının değişkenliği", 3)
    P_cite(doc, [
        "Yaşlılarda ateş eşiği ≥38,3°C veya bazal sıcaklıktan ≥1,3°C artış olarak kabul edilir",
        2, ". Ciddi bakteriyel enfeksiyonlarda ateşsiz seyir (%50–65) sık görülür; hipotermi kötü prognoz işareti olabilir ",
        7, ". Ateş tek başına tanı veya tedavi kriteri değildir."
    ])

    caption(doc, 2, "Geriatrik sendromlar ve olası enfeksiyon odağı")
    table(doc, ["Atipik bulgu", "Olası odak", "Dikkat"], [
        ["Yeni deliryum", "İYE, pnömoni, sepsis", "Tek başına İYE tanısı için yeterli değil"],
        ["Düşme / immobilite", "Sistemik enfeksiyon", "Mekanik düşme sanılabilir"],
        ["İştahsızlık, halsizlik", "Solunum yolu, İYE, sepsis", "Depresyon sanılmamalı"],
        ["İdrar inkontinansı", "İYE (destekleyici)", "Asimptomatik bakteriüri ile karışmamalı"],
        ["Kronik hastalık alevlenmesi", "Altta yatan enfeksiyon", "KOAH, KY, diyabet dekompansasyonu"],
    ], [4, 4.5, 6])

    H(doc, "Sık tanı hataları", 3)
    errors = [
        ("Asimptomatik bakteriüriyi İYE sanmak:", " Antibiyotik endikasyonu semptom gerektirir.", [9]),
        ("Deliryumu otomatik İYE kabul etmek:", " Mental durum değişikliği tek başına güvenilir değildir.", [8]),
        ("Pozitif kültürü enfeksiyon saymak:", " Kolonizasyon ve kontaminasyon ayırt edilmeli.", [1]),
        ("qSOFA’yı tek tarama aracı kullanmak:", " Yaşlıda duyarlılık sınırlıdır; NEWS/NEWS2 destekleyicidir.", [4]),
    ]
    for i, (a, b, refs) in enumerate(errors, 1):
        segs = [f"{i}. {a}{b}"]
        for r in refs:
            segs.append(r)
        P_cite(doc, segs, indent=False)

    H(doc, "TANISAL YAKLAŞIM: ALGORİTMİK ÇERÇEVE")
    P_cite(doc, ["Yaşlıda enfeksiyon tanısı, “enfeksiyon var mı?” ile “hangi odak?” sorularının ayrı adımlarda yanıtlanmasını gerektirir ", 1, 3, "."])

    H(doc, "Adım 1: Baseline değişiklik var mı?", 3)
    P(doc, "Son 24–72 saatte vital bulgu değişimi, yeni deliryum, fonksiyonel gerileme, düşme, inkontinans artışı veya kronik hastalıkta açıklanamayan kötüleşme varsa yüksek indeksli değerlendirme başlatılır.")

    H(doc, "Adım 2: Aciliyet sınıflaması", 3)
    add_img(doc, figs["s2"], 2, "Tanısal yaklaşım: aciliyet sınıflaması algoritması")
    P(doc, "Septik şok veya hemodinamik instabilitede acil stabilizasyon, kan/kültür, görüntüleme ve 1 saat içinde empirik geniş spektrum antibiyotik başlanır. Orta-ağır sistemik bulguda (NEWS/NEWS2 yükselmiş, laktat artmış, AKI) hastane yatışı değerlendirilir; hedefe yönelik örnek alınır; empirik tedavi 3 saat içinde başlatılır. Stabil ve lokalize bulgularda ayaktan veya yarı yatarak yönetim ile sistematik odak taraması yapılır.")

    H(doc, "Adım 3: Odak arama (bedside sistematik tarama)", 3)
    systems = "Baş–boyun (oral/diş, fasiyal selülit) | Solunum (taşipne, hipoksemi, ral) | Kardiyovasküler (yeni üfürüm, perfüzyon) | Abdomen (divertikülit, kolitis) | Genitoüriner (dizüri, retansiyon) | Deri-yumuşak doku (dekubitus, selülit, diyabetik ayak) | Kateter/cihaz | Sinir sistemi (menenjit bulguları, atipik olabilir)."
    P(doc, systems)

    H(doc, "Adım 4: Laboratuvar ve görüntüleme", 3)
    P_cite(doc, [
        "Temel panel (endikasyona göre): tam kan sayımı, CRP/prokalsitonin (yorum dikkatli), böbrek/karaciğer fonksiyonları, elektrolitler, laktat (sepsis şüphesi), idrar analizi. Kültür prensibi: tedavi başlamadan önce kan kültürü (sepsis şüphesi); odak odaklı örnekler (idrar, balgam, yara aspirasyonu). Kontaminasyonu azaltmak için uygun alım tekniği şarttır. İYE tanısında: lokalizasyon bulgusu → piyüri → kültür. Piyüri (−) ise İYE olasılığı düşüktür (NPV >%95) ",
        1, 3, ". Görüntüleme: pnömoni şüphesinde göğüs grafisi; komplike İYE/abdominal odakta USG/BT; endokardit şüphesinde ekokardiyografi."
    ])

    H(doc, "Adım 5: Alternatif tanıları dışla", 3)
    P(doc, "Metabolik bozukluk, ilaç yan etkisi, akut koroner sendrom, inme, subdural hematom, idrar retansiyonu, kabızlık, ağrı ve depresyon ayırıcı tanıda düşünülmelidir.")

    H(doc, "TEDAVİ YAKLAŞIMI: ALGORİTMİK VE BİREYSELLEŞTİRİLMİŞ YÖNETİM")
    P(doc, "Tedavi kararı dört eksende verilir: (1) Enfeksiyon ciddiyeti ve odak, (2) Hasta kırılganlığı, (3) Komorbid hastalıklar ve polifarmasi, (4) Hasta hedefleri / bakım planı (ACP).")
    P_cite(doc, ["Yaşlıda empirik tedavi standart rehber dozu ile başlamaz; önce komorbidite filtresi uygulanır ", 11, 12, "."])

    H(doc, "Aşama 0: Tedavi öncesi zorunlu tarama", 3)
    caption(doc, 3, "Tedavi öncesi komorbidite ve ilaç tarama kontrol listesi")
    table(doc, ["Alan", "Bakılacak", "Tedaviyi nasıl değiştirir?"], [
        ["Böbrek fonksiyonu", "eGFR, idrar çıkışı, diyaliz", "Doz ayarı; nefrotoksik ilaçlardan kaçın"],
        ["Kalp yetmezliği", "EF, volüm, BNP", "Sıvı stratejisi; kardiyotoksik ilaçlardan kaçın"],
        ["KOAH / astım", "O₂, steroid", "Solunum desteği; makrolid etkileşimi"],
        ["Diyabet", "HbA1c, OAD", "Glisemi; SGLT2i ile euglise riski"],
        ["Karaciğer", "Child-Pugh", "Hepatotoksik ilaçlardan kaçın"],
        ["Antikoagülan", "Varfarin, DOAC", "TMP-SMX, makrolid etkileşimi"],
        ["QT uzaması", "EKG", "Makrolid, kinolon, azol dikkat"],
        ["CDI öyküsü", "Son kolit", "4C antibiyotiklerden kaçın"],
        ["Demans/deliryum", "4AT", "Deliryum önleyici bakım"],
        ["İmmünosupresyon", "Steroid, biyolojik", "Genişletilmiş empirisi"],
        ["Prostat/BPH", "Retansiyon", "İdrar boşaltımı öncelikli"],
        ["Polifarmasi", "≥5 ilaç", "Etkileşim taraması"],
    ], [3.2, 4.3, 6])

    add_img(doc, figs["s3"], 3, "Antibiyotik başlama kararı algoritması (Aşama 1)")

    algo_box(doc, [
        "Enfeksiyon olasılığı değerlendirildi",
        "  DÜŞÜK → Antibiyotik BAŞLAMA; alternatif nedenleri tedavi et; 24–48 saatte yeniden değerlendir",
        "  ORTA + stabil → Aşama 0 tamamla; kültür al; Aşama 2+3; 48–72 saatte de-eskale",
        "  YÜKSEK / SEPSİS → Acil stabilizasyon; kan kültürü; geniş spektrum ≤1 sa (şok); de-eskale",
    ], 10, "Antibiyotik başlama kararı – metin algoritması (Aşama 1)")

    P_cite(doc, [
        "Niçin bireyselleştirme? Kırılgan yaşlıda tedavi etmeme deliryum/sepsis riski taşır; standart reçete ise C. difficile, AKI, aritmi ve kanama riskini artırır ",
        11, 12, "."
    ])

    H(doc, "Aşama 2: Enfeksiyon odağına göre empirik seçim", 3)
    P(doc, "Bu aşamada genel rehber seçimi yapılır; bir sonraki aşamada komorbiditeye göre değiştirilir.")
    add_img(doc, figs["s6"], 4, "Enfeksiyon odağına göre empirik antibiyotik seçimi (Aşama 2)")

    odak_detay = [
        "Üst solunum yolu (viral olasılık yüksek, stabil): Semptomatik tedavi; antibiyotik genellikle gerekmez. İstisna: bakteriyel süperenfeksiyon, KOAH alevlenmesi.",
        "Alt solunum yolu / pnömoni: Toplum kökenli → amoksisilin veya respiratuvar kinolon (KY/KBH/QT varsa değiştir). Hastane/UDBK kökenli → anti-pseudomonal β-laktam ± vankomisin.",
        "İYE (semptomatik, komplike değil): TMP-SMX veya nitrofurantoin (eGFR uygunsa) veya fosfomisin. eGFR <30: nitrofurantoin kullanılmaz.",
        "Komplike İYE / piyelonefrit: Sefalosporin veya piperasilin-tazobaktam. CDI öyküsünde sefalosporin değiştirilir.",
        "Deri-yumuşak doku: Sefazolin/dikloksasilin; MRSA riskinde TMP-SMX, doksisiklin veya vankomisin.",
        "İntra-abdominal / divertikülit: Metronidazol + sefalosporin veya piperasilin-tazobaktam.",
        "Sepsis (odak belirsiz): Meropenem veya piperasilin-tazobaktam + vankomisin; kaynak kontrolü planlanır.",
    ]
    for item in odak_detay:
        P(doc, "• " + item, indent=False)

    P(doc, "Genel ilkeler (tüm odaklar için): Yerel antibiyograma uyum; dar spektrum ve kısa süre (5–7 gün); 48–72 saatte de-eskale; polifarmasi etkileşim kontrolü.")

    H(doc, "Aşama 3: Komorbidite filtresi — standart reçeteyi değiştir", 3)
    P(doc, "Empirik seçim yapıldıktan sonra her komorbidite için aşağıdaki dallanma uygulanır. Birden fazla komorbidite varsa en kısıtlayıcı kural önceliklidir.")

    H(doc, "3A. Böbrek yetmezliği / KBH", 3)
    P_cite(doc, [
        "eGFR 30–59 arasında vankomisin, acyclovir ve β-laktam dozları azaltılmalı; nitrofurantoin ve aminoglikozidlerden kaçınılmalıdır. eGFR <30 veya diyalizde nitrofurantoin kullanılmaz; TMP-SMX genellikle kaçınılır (hiperkalemi, kemik iliği baskılanması). Kinolon doz ayarı yapılır ancak QT/KY birlikteliğinde tercih edilmez. Geçici AKI’de nefrotoksikler durdurulur, volüm resüsitasyonu sonrası doz yeniden hesaplanır ",
        11, 12, ". Tercih: ceftriaxone, amoksisilin-klavulanat, meropenem, linezolid."
    ])
    algo_box(doc, [
        "KBH (eGFR düşük / diyaliz):",
        "  eGFR 30–59: doz azalt; nitrofurantoin KAÇIN; aminoglikozid KAÇIN",
        "  eGFR <30: nitrofurantoin KULLANMA; TMP-SMX genellikle KAÇIN",
        "  Tercih: ceftriaxone, amoksisilin-klavulanat, meropenem, linezolid",
        "  Geçici AKI: nefrotoksikleri durdur; volüm sonrası dozu yeniden hesapla",
    ], 5, "Böbrek yetmezliğinde antibiyotik seçimi (Aşama 3A)")

    H(doc, "3B. Kalp yetmezliği", 3)
    P_cite(doc, [
        "KY’li yaşlıda pnömoni/sepsis mortalitesi yüksektir; fakat 30 mL/kg agresif kristalloid bolus genellikle uygun değildir. Küçük alıkverişli bolus (250–500 mL) ve dinamik yanıt (POCUS, BNP, JVP) önerilir. Hipotansiyonda erken noradrenalin; diüretik geçici kesilir. Makrolid ve kinolon QT riski nedeniyle kaçınılır ",
        4, 10, ". Aşırı sıvı pulmoner ödem, deliryum ve KY alevlenmesi yapabilir."
    ])
    algo_box(doc, [
        "KALP YETMEZLİĞİ:",
        "  Sepsis: 30 mL/kg bolus UYGUN DEĞİL → 250–500 mL + dinamik yanıt",
        "  Hipotansiyon: erken noradrenalin; diüretiği geçici kes",
        "  Antibiyotik: makrolid/kinolon QT → KAÇIN",
        "  Tercih: amoksisilin, sefalosporin, doksisiklin, linezolid",
    ], 6, "Kalp yetmezliğinde enfeksiyon tedavisi (Aşama 3B)")

    H(doc, "3C. KOAH / kronik solunum hastalığı", 3)
    P(doc, "Pnömoni/KOAH alevlenmesinde SpO₂ hedefi 88–92% olmalı; bronkodilatatör ve steroid protokolü uygulanır. Antibiyotik endikasyonu purulan balgam ve artan dispne ile birlikte değerlendirilir. Teofilin kullanan hastada makrolid/kinolon teofilin toksisitesi yapabilir; düzey izlenmelidir. Sepsiste non-invaziv ventilasyon erken düşünülmeli; sedasyon minimal tutulmalıdır (CO₂ retansiyonu, deliryum).")

    H(doc, "3D. Diyabet mellitus", 3)
    P(doc, "Her enfeksiyonda ayak muayenesi zorunludur; diyabetik ayak enfeksiyonu atlanmamalıdır. Enfeksiyon hiperglisemi yapar; glisemi sık izlenir. Sepsis, AKI veya kontrast varsa metformin geçici durdurulur (laktik asidoz riski). SGLT2 inhibitörleri euglise riski nedeniyle sepsis/cerrahi döneminde durdurulur. Kinolon hipoglisemi ve tendinit riski taşır; alternatif tercih edilir. Diyabetik ayak enfeksiyonunda iskemi değerlendirilir; MRSA + anaerob cover, cerrahi debridman ve vasküler konsültasyon gerekebilir.")

    H(doc, "3E. Karaciğer hastalığı / siroz", 3)
    P(doc, "Child-Pugh B–C hastada izoniazid, ketokonazol ve eritromisin hepatotoksisite riski nedeniyle kaçınılır. Sefalosporinler genellikle güvenlidir. Spontan bakteriyel peritonit şüphesinde sefalosporin 3. kuşak + albümin protokolü uygulanır. Sirozlu hastada sık varfarin kullanımı nedeniyle TMP-SMX, metronidazol ve makrolid INR artışına yol açabilir.")

    H(doc, "3F. Antikoagülan / kanama riski", 3)
    P_cite(doc, [
        "Varfarin ile TMP-SMX ciddi INR artışı yapar; alternatif tercih edilmeli veya sık INR izlenmelidir. Metronidazol ve makrolidler de etkileşim gösterir. DOAC kullanan hastada TMP-SMX, ketokonazol ve makrolid plazma düzeyini artırır. Linezolid trombositopeni yapabilir; düşük trombosit veya aktif kanamada kaçınılır ",
        12, 13, "."
    ])

    H(doc, "3G. QT uzaması", 3)
    P(doc, "EKG’de QTc >500 ms, amiodaron/sotalol kullanımı veya elektrolit bozukluğunda makrolid, kinolon ve azol antifungaller kaçınılır. Tercih: β-laktam, doksisiklin, linezolid. Mg/K düzeltilir; EKG tekrarlanır.")

    H(doc, "3H. C. difficile öyküsü", 3)
    P(doc, "Son antibiyotik, PPI ve hospitalizasyon öyküsü olan hastada 4C grubundan (co-amoksislav, kinolonlar, 3. kuşak sefalosporinler, klindamisin) kaçınılır. Tercih: amoksisilin (klavulanatsız), doksisiklin, linezolid. Oral vankomisin profilaksisi rutin değildir.")

    H(doc, "3I. Demans, deliryum, Parkinson", 3)
    P_cite(doc, [
        "Deliryum önleyici bakım paketi (görme/işitme desteği, mobilizasyon, uyku hijyeni) uygulanır. Kinolon deliryum riskini artırır. Parkinson hastasında metoklopramid ve proklorperazin kesinlikle kaçınılır; levodopa kesilmemelidir ",
        8, 11, "."
    ])

    H(doc, "3J. İmmünosupresyon", 3)
    P(doc, "Ateşsiz ciddi enfeksiyon olabilir; düşük eşikle değerlendirme yapılır. Genişletilmiş empirisi düşünülür (pseudomonas, fungus). Nötropenide febril nötropeni protokolü (piperasilin-tazobaktam ± vankomisin) uygulanır. Canlı aşılar tedavi süresince kontrendikedir.")

    H(doc, "3K. Prostat hipertrofisi / idrar retansiyonu", 3)
    P(doc, "Enfeksiyon tedavisinden önce idrar retansiyonu giderilmelidir (kateter/sonda). Antikolinergik ilaçlar geçici kesilir. Kalıcı kateter enfeksiyon kaynağı olabilir; çıkar/değiştir planı yapılır.")

    algo_box(doc, [
        "ÖZET — Diğer komorbiditeler (3C–3K):",
        "KOAH: SpO₂ 88–92%; NIV; teofilin etkileşimi",
        "DİYABET: ayak muayenesi; metformin/SGLT2i durdur",
        "ANTİKOAGÜLAN: TMP-SMX/makrolid → INR/kanama",
        "QT: makrolid, kinolon, azol KAÇIN",
        "CDI: 4C grubu KAÇIN | PARKİNSON: metoklopramid YASAK",
    ], 7, "Komorbidite modifikasyonları özet şeması (Aşama 3C–3K)")

    H(doc, "Aşama 4: Birleşik komorbidite algoritması", 3)
    P(doc, "Birden fazla komorbiditesi olan tipik yaşlı hasta için antibiyotik seçimi aşağıdaki adımlarla yürütülür:")
    for step in [
        "Adım 1 — Acil tehlike (septik şok, ciddi hipoksi, menenjit): Geniş spektrum hemen; komorbidite doz/ilaç düzeltmesi sonraki 24 saatte.",
        "Adım 2 — Mutlak kaçınılacaklar: CDI öyküsü → 4C çıkar; eGFR <30 → nitrofurantoin/aminoglikozid çıkar; QT → makrolid/kinolon çıkar; antikoagülan → TMP-SMX/makrolid dikkat; Parkinson → metoklopramid yasak.",
        "Adım 3 — Odak-spesifik tercih (Aşama 2).",
        "Adım 4 — Komorbidite modifikasyonu: KBH+İYE → fosfomisin/amoksisilin; KBH+sepsis → meropenem; KY+pnömoni → amoksisilin-klavulanat; KY+sepsis → küçük bolus + vazopressör; diyabet+ayak → MRSA + cerrahi; KOAH+pnömoni → β-laktam + O₂/NIV.",
        "Adım 5 — Rota ve süre: Oral geçiş afebrile 24–48 saatte; çoğu enfeksiyon 5–7 gün; 48–72 saatte klinik yanıt değerlendir.",
    ]:
        P(doc, step, indent=False)

    add_img(doc, figs["s4"], 8, "Birleşik komorbidite algoritması – pratik karar ağacı (Aşama 4)")

    caption(doc, 4, "Komorbiditeye göre kaçınılacak ve tercih edilecek antibiyotikler")
    table(doc, ["Komorbidite", "Kaçınılması", "Tercih alternatifleri"], [
        ["KBH (eGFR <30)", "Nitrofurantoin, aminoglikozid, TMP-SMX", "Amoksisilin, ceftriaxone, meropenem"],
        ["Kalp yetmezliği", "Agresif sıvı; QT uzatan ilaçlar", "β-laktam; küçük bolus + vazopressör"],
        ["KOAH", "Aşırı O₂", "Amoksisilin-klavulanat; NIV"],
        ["Diyabet", "Kinolon; metformin (sepsis/AKI)", "Sefalosporin; glisemi izlemi"],
        ["Varfarin/DOAC", "TMP-SMX, makrolid", "Amoksisilin, doksisiklin"],
        ["QT uzaması", "Makrolid, kinolon, azol", "β-laktam, linezolid"],
        ["CDI öyküsü", "4C grubu", "Amoksisilin, doksisiklin, linezolid"],
        ["Parkinson", "Metoklopramid, kinolon", "Amoksisilin, sefalosporin"],
        ["Demans", "Kinolon", "β-laktam; deliryum önleyici bakım"],
        ["İmmünosupresyon", "Dar spektrum gecikmesi", "Geniş spektrum empirisi"],
    ], [3.5, 5, 5.5])

    H(doc, "Aşama 5: Sepsis yönetimi — komorbiditeye göre uyarlanmış şema", 3)
    P_cite(doc, [
        "Surviving Sepsis Campaign 2021 erken antibiyotik ve kaynak kontrolünü vurgular; yaşlı kırılgan hastada frailty ve ACP ile birlikte invaziv destek kararı verilmelidir ",
        10, 4, 14, "."
    ])
    add_img(doc, figs["s5"], 9, "Sepsis yönetimi: komorbiditeye göre sıvı stratejisi (Aşama 5)")

    sepsis_detay = [
        "Antibiyotik: Septik şokta 1 saat içinde geniş spektrum (Aşama 3–4 filtresi ile doz/ilaç). Sepsis (şoksuz): en kısa sürede başla; 48–72 saatte de-eskale.",
        "Sıvı — KY yok, hipovolemik: kısıtlı bolus (500 mL) + yanıt değerlendir.",
        "Sıvı — KY var, EF düşük: 250 mL bolus + POCUS/JVP; yanıt yoksa vazopressör; diüretik geçici kes.",
        "Sıvı — KBH + diyaliz: volüm durumu diyaliz ekibi ile planlanır; nefrotoksiklerden kaçın.",
        "Vazopressör/YBÜ: Frailty yüksek + ACP konfor odaklı ise invaziv destek sınırlı olabilir; multidisipliner karar.",
        "Destekleyici: kaynak kontrol, deliryum önleme, erken mobilizasyon, palyatif entegrasyon.",
    ]
    for item in sepsis_detay:
        P(doc, "• " + item, indent=False)

    P(doc, "qSOFA yaşlıda duyarlılığı düşüktür; KY/KBH ile organ yetmezliği sınırları belirsizleşir. Klinisyen muayenesi + seri vital + laktat esastır.")

    H(doc, "Aşama 6: Polifarmasi ve ilaç–ilaç etkileşim kontrolü", 3)
    caption(doc, 5, "Antimikrobiyal ilaç–ilaç etkileşimleri (yaşlı hasta)")
    table(doc, ["Antibiyotik", "Etkileşim", "Sonuç", "Yapılacak"], [
        ["Klarytromisin", "Statin, varfarin", "Rabdomiyoliz, kanama", "Alternatif AB"],
        ["TMP-SMX", "Varfarin, ACEi", "INR↑, hiperkalemi", "K/INR izle"],
        ["Metronidazol", "Varfarin", "INR↑", "Alternatif"],
        ["Kinolon", "Teofilin, OAD", "Toksisite, hipoglisemi", "Alternatif"],
        ["Vankomisin", "Aminoglikozid", "Nefrotoksisite", "Düzey izle"],
        ["Meropenem", "Valproat", "Düzey değişimi", "AED izle"],
    ], [2.8, 3.2, 3.5, 4])

    P_cite(doc, ["Polifarmasi yönetiminde her yeni antibiyotik için etkileşim taraması zorunludur ", 13, "."])

    H(doc, "Kaynak kontrolü ve bakım desteği", 3)
    for item in [
        "Kateter çıkarılması/değişimi (idrar, santral venöz)",
        "İdrar retansiyonu giderilmesi (BPH, nörogen mesane)",
        "Dekubitus bakımı ve bası yarası debridmanı",
        "Aspirasyon önleme; yutma değerlendirmesi",
        "KY/KBH’ye uygun hidratasyon; erken mobilizasyon",
        "Glisemi, elektrolit, böbrek fonksiyonu günlük izlem",
    ]:
        P(doc, "• " + item, indent=False)

    H(doc, "ÖNLEME, AŞILAMA VE ENFEKSİYON KONTROLÜ")
    P_cite(doc, [
        "Yaşlıda aşı yanıtı azalmış olsa da influenza, pnömokok, COVID-19, herpes zoster ve RSV aşıları birincil korunmanın temelidir ",
        15, "."
    ])
    for v in ["Influenza: yıllık, yüksek doz/adjuvanlı formlar.", "Pnömokok: konjuge aşı serisi.", "COVID-19: güncel takviye doz.", "Herpes zoster: rekombinant aşı.", "RSV: yaşlı popülasyonda yeni seçenekler."]:
        P(doc, "• " + v, indent=False)
    P_cite(doc, ["UDBK’larda 2025 multisociety enfeksiyon kontrol rehberi el hijyeni, izolasyon, personel eğitimi, aşı programları, cihaz ilişkili enfeksiyon paketleri ve salgın yönetimini vurgular ", 16, ". Yaşlı hasta popülasyonunun büyük bölümü UDBK’da izlendiğinden, ayaktan hekimin de bu protokolleri bilmesi gerekir."])
    P(doc, "Antimikrobiyal stewardship (AMS) hedefleri:")
    for ams in [
        "Gereksiz antibiyotik azaltımı (asimptomatik bakteriüri, viral üSYE)",
        "Kısa etkin süre (çoğu enfeksiyon 5–7 gün)",
        "Dar spektrum tercih",
        "Klinik yanıt odaklı erken durdurma",
        "“Non-lokalize bulgu = antibiyotik” refleksinden kaçınma",
    ]:
        P(doc, "• " + ams, indent=False)
    P_cite(doc, [
        "Niçin aşılama bu bölümde? Enfeksiyon yönetimi yalnızca tedavi değil; hastane yatışını, deliryumu ve antibiyotik maruziyetini azaltan üst düzey stratejidir ",
        15, "."
    ])

    H(doc, "PRATİK KLİNİK SENARYOLAR")
    scenarios = [
        ("Senaryo 1 — 84 yaş, UDBK, yeni konfüzyon", "Yanlış: Otomatik İYE + kinolon. Doğru: 4AT, vital, odak taraması, idrar tahlili (piyüri/kültür destekleyici bulgularla), göğüs grafisi; alternatif nedenler (retansiyon, ilaç, metabolik). Antibiyotik yalnızca yüksek olasılıkta.", [8]),
        ("Senaryo 2 — 78 yaş, diyabet, ayak kızarıklığı", "Yanlış: Oral sefalosporin + ayakta izlem. Doğru: nabız/iskemi, derin doku değerlendirmesi, MRSA cover, cerrahi debridman; glisemi optimizasyonu.", [12]),
        ("Senaryo 3 — 90 yaş, kırılgan, sepsis", "Yanlış: Gecikme. Doğru: Sepsis protokolü + frailty/ACP; erken antibiyotik; multidisipliner karar.", [4, 14]),
        ("Senaryo 4 — 82 yaş, KY+KBH, pnömoni+hipotansiyon", "Yanlış: 30 mL/kg sıvı + levofloksasin + TMP-SMX. Doğru: 250 mL bolus + POCUS; noradrenalin; amoksisilin-klavulanat/meropenem (doz ayarlı); diüretik kes; günlük K/kreatinin.", [10, 11]),
        ("Senaryo 5 — 76 yaş, varfarin+Parkinson, İYE", "Yanlış: TMP-SMX + metoklopramid. Doğru: fosfomisin/amoksisilin; INR izlemi; metoklopramid yasak.", [11, 13]),
    ]
    for title, body, refs in scenarios:
        P(doc, title, bold=True, indent=False)
        if "Yanlış:" in body and "Doğru:" in body:
            parts = body.split("Doğru:")
            P(doc, parts[0].strip(), indent=True)
            segs = ["Doğru: " + parts[1].strip()]
            for r in refs:
                segs.append(r)
            P_cite(doc, segs)
        else:
            segs = [body]
            for r in refs:
                segs.append(r)
            P_cite(doc, segs)

    P(doc, "Bu senaryolar, komorbidite filtresinin klinik uygulamada neden kritik olduğunu göstermektedir. Özellikle KY+KBH birlikteliğinde hem sıvı hem antibiyotik seçimi standart protokolden sapmalı; Parkinson ve antikoagülan kullanan hastada ise ilaç–ilaç etkileşimleri yaşamı tehdit edebilir.")

    H(doc, "SONUÇ")
    for i, item in enumerate([
        "Atipik sunumu tanımak (deliryum, düşme, fonksiyon kaybı)",
        "Sistematik odak araması yapmak",
        "Komorbidite ve polifarmasiyi tedavi kararına entegre etmek",
        "Erken uygun antimikrobiyal tedavi ile gereksiz kullanımı dengelemek",
        "Aşılama ve enfeksiyon kontrolü ile birincil önlem sağlamak",
    ], 1):
        P(doc, f"{i}. {item}", indent=False)
    P(doc, "Algoritmik yaklaşım klinik muhakemenin yerini almaz; gecikmeyi azaltır, tanı hatalarını sınırlar ve bireyselleştirilmiş bakım için çerçeve sunar.")

    H(doc, "KAYNAKLAR")
    for i, ref in enumerate(REFS, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        suffix = " (Köşe taşı)" if i in CORNERSTONE else ""
        r = p.add_run(f"{i}. {ref}{suffix}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    out = f"{OUT_DIR}/Yaslilarda-Enfeksiyon-Hastaliklarina-Genel-Yaklasim.docx"
    doc.save(out)

    words = sum(len(p.text.split()) for p in doc.paragraphs)
    for tb in doc.tables:
        for row in tb.rows:
            for c in row.cells:
                words += len(c.text.split())
    print(f"Saved: {out}")
    print(f"Words: ~{words}, Est. pages: {words/280:.1f}")
    return out


if __name__ == "__main__":
    build()
